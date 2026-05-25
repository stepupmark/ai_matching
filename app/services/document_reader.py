"""
Document reader service — extracts plain text from uploaded document files.

Supported formats:
  - PDF  (via PyMuPDF / fitz)
  - DOCX (via python-docx) — Open XML format only (.docx), not legacy binary .doc
  - TXT  (plain text)

Each extraction function accepts raw file bytes and returns the extracted text.
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Optional

from app.config import settings

logger = logging.getLogger(__name__)

# ── Supported extensions ──────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".text"}
# Note: Legacy binary .doc format is NOT supported.
# python-docx only handles .docx (Open XML). Convert .doc files to .docx first.

MAX_FILE_SIZE = settings.max_file_size_mb * 1024 * 1024


class DocumentReadError(Exception):
    """Raised when a document cannot be read or parsed."""


class UnsupportedFormatError(DocumentReadError):
    """Raised when the file format is not supported."""


def extract_text(content: bytes, filename: str) -> str:
    """
    Extract plain text from a document file.

    Auto-detects format from the file extension.

    Args:
        content: Raw file bytes.
        filename: Original filename (used to detect format).

    Returns:
        Extracted plain text.

    Raises:
        UnsupportedFormatError: If the file extension is not supported.
        DocumentReadError: If the document cannot be parsed.
    """
    if len(content) > MAX_FILE_SIZE:
        raise DocumentReadError(
            f"File exceeds maximum size of {settings.max_file_size_mb} MB"
        )

    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFormatError(
            f"Unsupported file format '{ext}'. Supported: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext == ".pdf":
        return _extract_pdf(content)
    elif ext == ".docx":
        return _extract_docx(content)
    elif ext in (".txt", ".text"):
        return _extract_txt(content)
    else:
        raise UnsupportedFormatError(f"Unsupported format: {ext}")


# ── PDF extraction ────────────────────────────────────────────────

def _extract_pdf(content: bytes) -> str:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise DocumentReadError(
            "PyMuPDF (fitz) is required for PDF extraction. "
            "Install it with: pip install PyMuPDF"
        )

    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise DocumentReadError(f"Failed to open PDF: {exc}")

    pages: list[str] = []
    for page_num in range(len(doc)):
        try:
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                pages.append(text)
        except Exception as exc:
            logger.warning("Failed to extract text from PDF page %d: %s", page_num, exc)
            continue

    doc.close()

    if not pages:
        raise DocumentReadError(
            "No extractable text found in PDF. "
            "The document may be scanned/image-based (OCR not supported)."
        )

    return _clean_text("\n\n".join(pages))


# ── DOCX extraction ───────────────────────────────────────────────

def _extract_docx(content: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
    except ImportError:
        raise DocumentReadError(
            "python-docx is required for DOCX extraction. "
            "Install it with: pip install python-docx"
        )

    try:
        file_stream = io.BytesIO(content)
        document = docx.Document(file_stream)
    except Exception as exc:
        raise DocumentReadError(f"Failed to open DOCX: {exc}")

    paragraphs: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in document.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    if not paragraphs:
        raise DocumentReadError("No extractable text found in DOCX.")

    return _clean_text("\n".join(paragraphs))


# ── TXT extraction ────────────────────────────────────────────────

def _extract_txt(content: bytes) -> str:
    """Extract text from a plain-text file (auto-detect encoding)."""
    encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
    for enc in encodings:
        try:
            text = content.decode(enc)
            return _clean_text(text)
        except (UnicodeDecodeError, UnicodeError):
            continue

    # Last resort: decode with errors='replace'
    text = content.decode("utf-8", errors="replace")
    return _clean_text(text)


# ── Text cleaning ─────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Clean extracted text: normalize whitespace, remove null bytes."""
    text = text.replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()
    return text
